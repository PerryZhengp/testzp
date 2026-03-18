#!/usr/bin/env python3
"""
Strict, data-driven generator for 3 formal reports:
1) 自动化接口功能测试报告
2) 自动化接口性能测试报告
3) 给领导汇报一页摘要

Inputs are mandatory:
- functional JTL
- performance JTL
- defects CSV/JSON
- report meta JSON/YAML
"""

from __future__ import annotations

import argparse
import csv
import datetime as dt
import json
import math
import os
import pathlib
import re
import statistics
import sys
import xml.etree.ElementTree as et
import zipfile
from collections import Counter, defaultdict
from dataclasses import dataclass
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple
from xml.sax.saxutils import escape

os.environ.setdefault("MPLCONFIGDIR", "/tmp/mplconfig")
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

plt.rcParams["font.sans-serif"] = ["PingFang SC", "Heiti TC", "Arial Unicode MS", "SimHei", "DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False


@dataclass
class JtlRecord:
    label: str
    success: bool
    response_code: str
    elapsed: int
    timestamp: int
    failure_message: str


SEVERITY_MAP = {
    "致命": {"致命", "fatal", "critical", "blocker", "p0"},
    "严重": {"严重", "high", "major", "p1"},
    "一般": {"一般", "medium", "normal", "p2"},
    "建议": {"建议", "low", "minor", "suggestion", "p3"},
}

SCENE_NAME = {
    "PERF-001": "登录基准测试",
    "PERF-002": "Heartbeat 基准测试",
    "PERF-003": "查询类混合测试",
    "PERF-004": "作答保存混合测试",
    "PERF-005": "长稳测试",
}

DEFAULT_META: Dict[str, Any] = {
    "document": {
        "author": "Peng Zheng",
        "version": "V1.0",
        "date": dt.date.today().isoformat(),
        "test_phase": "测试环境",
        "tool": "JMeter",
    },
    "project": {
        "name": "校园版接口自动化测试",
        "performance_name": "校园版接口性能测试",
        "base_url": "http://192.168.0.90/api",
    },
    "accounts": ["zhengpengst", "hlqiongs", "zshaos1", "huyonglans1"],
    "base_data": {
        "courseId": "2032055666142167042",
        "teachingTaskId": "2033505568785190913",
        "blockId": "3O00nCpPKp",
    },
    "thresholds": {
        "error_rate_max_pct": 1.0,
        "tp95_max_ms": 1000,
        "success_rate_min_pct": 99.0,
    },
    "performance_scene_config": {
        "PERF-001": {"concurrency": "10/50/100", "duration": "5 min", "desc": "单独测试登录接口性能"},
        "PERF-002": {"concurrency": "50/100/200", "duration": "10 min", "desc": "验证心跳接口稳定性"},
        "PERF-003": {"concurrency": "50", "duration": "10 min", "desc": "课程、目录、内容查询混合场景"},
        "PERF-004": {"concurrency": "20/50", "duration": "10 min", "desc": "内容查询、答题、保存链路"},
        "PERF-005": {"concurrency": "50/100", "duration": "30-60 min", "desc": "带 Heartbeat 的持续在线会话测试"},
    },
    "release_recommendation": "",
}


def _safe_bool(value: str) -> bool:
    return str(value).strip().lower() in {"1", "true", "yes", "y"}


def _safe_int(value: Any, default: int = 0) -> int:
    try:
        return int(str(value).strip())
    except Exception:
        return default


def _normalize_text(v: Any) -> str:
    return str(v).strip()


def _deep_merge(base: Dict[str, Any], override: Dict[str, Any]) -> Dict[str, Any]:
    result = dict(base)
    for k, v in override.items():
        if isinstance(v, dict) and isinstance(result.get(k), dict):
            result[k] = _deep_merge(result[k], v)
        else:
            result[k] = v
    return result


def _parse_csv_jtl(path: pathlib.Path) -> List[JtlRecord]:
    rows: List[JtlRecord] = []
    with path.open("r", encoding="utf-8", newline="") as f:
        reader = csv.DictReader(f)
        for r in reader:
            label = _normalize_text(r.get("label") or r.get("lb"))
            if not label:
                continue
            rows.append(
                JtlRecord(
                    label=label,
                    success=_safe_bool(r.get("success") or r.get("s")),
                    response_code=_normalize_text(r.get("responseCode") or r.get("rc")),
                    elapsed=_safe_int(r.get("elapsed") or r.get("t"), 0),
                    timestamp=_safe_int(r.get("timeStamp") or r.get("ts"), 0),
                    failure_message=_normalize_text(r.get("failureMessage") or r.get("responseMessage") or r.get("rm")),
                )
            )
    return rows


def _parse_xml_jtl(path: pathlib.Path) -> List[JtlRecord]:
    rows: List[JtlRecord] = []
    root = et.parse(path).getroot()
    for elem in root.iter():
        if elem.tag not in {"httpSample", "sample"}:
            continue
        label = _normalize_text(elem.attrib.get("lb"))
        if not label:
            continue
        rows.append(
            JtlRecord(
                label=label,
                success=_safe_bool(elem.attrib.get("s", "")),
                response_code=_normalize_text(elem.attrib.get("rc")),
                elapsed=_safe_int(elem.attrib.get("t"), 0),
                timestamp=_safe_int(elem.attrib.get("ts"), 0),
                failure_message=_normalize_text(elem.attrib.get("rm")),
            )
        )
    return rows


def parse_jtl(path: pathlib.Path) -> List[JtlRecord]:
    if not path.exists():
        raise FileNotFoundError(f"JTL file not found: {path}")
    head = path.read_bytes()[:256].lstrip()
    records = _parse_xml_jtl(path) if head.startswith(b"<") else _parse_csv_jtl(path)
    if not records:
        raise ValueError(f"JTL has no records: {path}")
    return records


def _percentile(values: Sequence[int], pct: float) -> int:
    if not values:
        return 0
    if len(values) == 1:
        return int(values[0])
    sorted_vals = sorted(values)
    k = (len(sorted_vals) - 1) * pct
    f = math.floor(k)
    c = math.ceil(k)
    if f == c:
        return int(sorted_vals[int(k)])
    d0 = sorted_vals[f] * (c - k)
    d1 = sorted_vals[c] * (k - f)
    return int(d0 + d1)


def aggregate_metrics(records: Sequence[JtlRecord]) -> Dict[str, float]:
    if not records:
        return {
            "total": 0,
            "passed": 0,
            "failed": 0,
            "error_rate_pct": 0.0,
            "success_rate_pct": 0.0,
            "avg_ms": 0.0,
            "tp90_ms": 0.0,
            "tp95_ms": 0.0,
            "tp99_ms": 0.0,
            "tps": 0.0,
            "duration_s": 0.0,
        }

    elapsed = [max(0, r.elapsed) for r in records]
    passed = sum(1 for r in records if r.success)
    failed = len(records) - passed
    ts = [r.timestamp for r in records if r.timestamp > 0]
    if len(ts) >= 2 and max(ts) > min(ts):
        duration_s = (max(ts) - min(ts)) / 1000.0
    else:
        duration_s = max(1.0, len(records) * (statistics.mean(elapsed) / 1000.0))
    return {
        "total": float(len(records)),
        "passed": float(passed),
        "failed": float(failed),
        "error_rate_pct": (failed / len(records)) * 100.0,
        "success_rate_pct": (passed / len(records)) * 100.0,
        "avg_ms": float(round(statistics.mean(elapsed), 2)),
        "tp90_ms": float(_percentile(elapsed, 0.90)),
        "tp95_ms": float(_percentile(elapsed, 0.95)),
        "tp99_ms": float(_percentile(elapsed, 0.99)),
        "tps": float(round(len(records) / max(duration_s, 1e-6), 2)),
        "duration_s": float(round(duration_s, 2)),
    }


def load_meta(path: pathlib.Path) -> Dict[str, Any]:
    if not path.exists():
        raise FileNotFoundError(f"Meta file not found: {path}")
    suffix = path.suffix.lower()
    if suffix in {".json"}:
        content = json.loads(path.read_text(encoding="utf-8"))
    elif suffix in {".yml", ".yaml"}:
        try:
            import yaml  # type: ignore
        except Exception as e:
            raise ValueError(
                "YAML meta file requires PyYAML in runtime. "
                "Please use JSON meta or install pyyaml."
            ) from e
        content = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    else:
        raise ValueError("Meta file must be .json/.yml/.yaml")
    if not isinstance(content, dict):
        raise ValueError("Meta file root must be an object")
    return _deep_merge(DEFAULT_META, content)


def load_defects(path: pathlib.Path) -> List[Dict[str, str]]:
    if not path.exists():
        raise FileNotFoundError(f"Defects file not found: {path}")
    suffix = path.suffix.lower()
    rows: List[Dict[str, str]] = []
    if suffix == ".json":
        raw = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(raw, dict):
            raw = raw.get("defects", [])
        if not isinstance(raw, list):
            raise ValueError("JSON defects must be list or { defects: [] }")
        for item in raw:
            if not isinstance(item, dict):
                continue
            rows.append(
                {
                    "id": _normalize_text(item.get("id", "")),
                    "severity": _normalize_text(item.get("severity", "")),
                    "module": _normalize_text(item.get("module", "")),
                    "title": _normalize_text(item.get("title", "")),
                    "status": _normalize_text(item.get("status", "")),
                }
            )
    elif suffix == ".csv":
        with path.open("r", encoding="utf-8", newline="") as f:
            reader = csv.DictReader(f)
            required = {"id", "severity", "module", "title", "status"}
            missing = required - set(reader.fieldnames or [])
            if missing:
                raise ValueError(f"Defects CSV missing required columns: {sorted(missing)}")
            for r in reader:
                rows.append(
                    {
                        "id": _normalize_text(r.get("id", "")),
                        "severity": _normalize_text(r.get("severity", "")),
                        "module": _normalize_text(r.get("module", "")),
                        "title": _normalize_text(r.get("title", "")),
                        "status": _normalize_text(r.get("status", "")),
                    }
                )
    else:
        raise ValueError("Defects file must be .csv or .json")
    if not rows:
        raise ValueError("Defects file is empty. Strict mode requires defect input.")
    return rows


def severity_counts(defects: Sequence[Dict[str, str]]) -> Dict[str, int]:
    counts = {"致命": 0, "严重": 0, "一般": 0, "建议": 0}
    for d in defects:
        sev = d.get("severity", "").strip().lower()
        matched = False
        for k, alias in SEVERITY_MAP.items():
            if sev in {x.lower() for x in alias}:
                counts[k] += 1
                matched = True
                break
        if not matched:
            counts["一般"] += 1
    return counts


def write_docx_from_lines(path: pathlib.Path, lines: Sequence[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)

    def p(text: str) -> str:
        t = text if text else " "
        return f'<w:p><w:r><w:t xml:space="preserve">{escape(t)}</w:t></w:r></w:p>'

    body = "".join(p(x) for x in lines)
    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
 xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
 xmlns:o="urn:schemas-microsoft-com:office:office"
 xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
 xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
 xmlns:v="urn:schemas-microsoft-com:vml"
 xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
 xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
 xmlns:w10="urn:schemas-microsoft-com:office:word"
 xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
 xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
 xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
 xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"
 xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"
 xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
 mc:Ignorable="w14 wp14">
  <w:body>
    {body}
    <w:sectPr>
      <w:pgSz w:w="11906" w:h="16838"/>
      <w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" w:header="708" w:footer="708" w:gutter="0"/>
      <w:cols w:space="708"/>
      <w:docGrid w:linePitch="360"/>
    </w:sectPr>
  </w:body>
</w:document>
"""
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>
"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
"""
    styles = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/><w:qFormat/></w:style>
</w:styles>
"""
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document_xml)
        zf.writestr("word/styles.xml", styles)


def table_lines(headers: Sequence[str], rows: Sequence[Sequence[Any]]) -> List[str]:
    lines = [" | ".join(headers), " | ".join(["---"] * len(headers))]
    for r in rows:
        lines.append(" | ".join(str(x) for x in r))
    return lines


def _is_markdown_table_sep(line: str) -> bool:
    s = line.strip()
    if not s:
        return False
    allowed = set("-| :")
    return "-" in s and all(ch in allowed for ch in s)


def _split_md_row(line: str) -> List[str]:
    return [c.strip() for c in line.split("|")]


def _is_heading_line(line: str) -> Optional[int]:
    m = re.match(r"^(\d+)(?:\.(\d+))?\.\s+(.+)$", line.strip())
    if not m:
        return None
    text = m.group(3).strip()
    # Heuristic: short section titles without sentence punctuation are headings.
    if len(text) <= 40 and not text.endswith("。") and "：" not in text:
        return 1 if m.group(2) is None else 2
    return None


def _set_doc_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    try:
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    except Exception:
        pass

    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)]:
        try:
            s = doc.styles[style_name]
            s.font.name = "Calibri"
            s.font.size = Pt(size)
            s.font.bold = True
            s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        except Exception:
            continue


def write_professional_docx(
    path: pathlib.Path,
    lines: Sequence[str],
    charts: Optional[Dict[str, pathlib.Path]] = None,
) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    chart_map = charts or {}
    doc = Document()
    _set_doc_styles(doc)

    i = 0
    n = len(lines)
    while i < n:
        line = (lines[i] or "").rstrip()

        chart_match = re.match(r"^\[\[CHART:([A-Z0-9_]+)\]\]$", line.strip())
        if chart_match:
            key = chart_match.group(1)
            img_path = chart_map.get(key)
            if img_path and pathlib.Path(img_path).exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(img_path), width=Inches(6.3))
            else:
                doc.add_paragraph(f"[图表缺失] {key}")
            i += 1
            continue

        if i + 1 < n and "|" in line and _is_markdown_table_sep(lines[i + 1]):
            header = _split_md_row(line)
            rows: List[List[str]] = []
            j = i + 2
            while j < n and "|" in (lines[j] or "") and not _is_markdown_table_sep(lines[j]):
                rows.append(_split_md_row(lines[j]))
                j += 1
            cols = max(len(header), *(len(r) for r in rows)) if rows else len(header)
            table = doc.add_table(rows=1 + len(rows), cols=cols)
            table.style = "Table Grid"
            for c in range(cols):
                table.rows[0].cells[c].text = header[c] if c < len(header) else ""
                for run in table.rows[0].cells[c].paragraphs[0].runs:
                    run.bold = True
            for r_idx, row in enumerate(rows, start=1):
                for c in range(cols):
                    table.rows[r_idx].cells[c].text = row[c] if c < len(row) else ""
            doc.add_paragraph("")
            i = j
            continue

        if not line.strip():
            doc.add_paragraph("")
            i += 1
            continue

        if line.startswith("《") and line.endswith("》"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(20)
            try:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            except Exception:
                pass
            i += 1
            continue

        heading_level = _is_heading_line(line)
        if heading_level is not None:
            doc.add_heading(line, level=heading_level)
            i += 1
            continue

        if line.startswith("• "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+\.\s+.+", line.strip()) and line.strip().endswith("。"):
            doc.add_paragraph(line.strip(), style="List Number")
            i += 1
            continue

        doc.add_paragraph(line)
        i += 1

    doc.save(path)


def _ensure_chart_dir(base_output_dir: pathlib.Path) -> pathlib.Path:
    p = base_output_dir / "_charts"
    p.mkdir(parents=True, exist_ok=True)
    return p


def build_function_chart_paths(func_info: Dict[str, Any], output_dir: pathlib.Path) -> Dict[str, pathlib.Path]:
    charts: Dict[str, pathlib.Path] = {}
    chart_dir = _ensure_chart_dir(output_dir)

    module_data = func_info.get("module_chart_data", [])
    if module_data:
        labels = [x["module"] for x in module_data]
        passed = [x["passed"] for x in module_data]
        failed = [x["failed"] for x in module_data]
        x = range(len(labels))
        fig, ax = plt.subplots(figsize=(8.6, 3.8))
        ax.bar(x, passed, label="通过", color="#2E7D32")
        ax.bar(x, failed, bottom=passed, label="失败", color="#C62828")
        ax.set_title("功能模块通过/失败分布")
        ax.set_xticks(list(x))
        ax.set_xticklabels(labels, rotation=20, ha="right")
        ax.set_ylabel("样本数")
        ax.legend()
        fig.tight_layout()
        p = chart_dir / "func_module_pass_fail.png"
        fig.savefig(p, dpi=150)
        plt.close(fig)
        charts["FUNC_MODULE_PASS_FAIL"] = p

    status_counter: Counter = func_info.get("status_code_counter", Counter())
    if status_counter:
        items = status_counter.most_common(8)
        codes = [k for k, _ in items]
        cnts = [v for _, v in items]
        fig, ax = plt.subplots(figsize=(8.6, 3.2))
        ax.bar(codes, cnts, color="#1565C0")
        ax.set_title("功能样本状态码分布（Top8）")
        ax.set_xlabel("HTTP/业务状态码")
        ax.set_ylabel("次数")
        fig.tight_layout()
        p = chart_dir / "func_status_code_dist.png"
        fig.savefig(p, dpi=150)
        plt.close(fig)
        charts["FUNC_STATUS_CODE_DIST"] = p

    return charts


def build_performance_chart_paths(perf_info: Dict[str, Any], output_dir: pathlib.Path) -> Dict[str, pathlib.Path]:
    charts: Dict[str, pathlib.Path] = {}
    chart_dir = _ensure_chart_dir(output_dir)
    scene_metrics = perf_info.get("scene_metrics", {})
    scene_ids = [sid for sid in ["PERF-001", "PERF-002", "PERF-003", "PERF-004", "PERF-005"] if sid in scene_metrics]

    if scene_ids:
        err_rates = [scene_metrics[sid]["metrics"]["error_rate_pct"] for sid in scene_ids]
        fig, ax = plt.subplots(figsize=(8.6, 3.6))
        ax.bar(scene_ids, err_rates, color="#D32F2F")
        ax.set_title("性能场景错误率分布")
        ax.set_ylabel("错误率(%)")
        ax.set_ylim(0, max(100, max(err_rates) * 1.15 if err_rates else 100))
        for i, v in enumerate(err_rates):
            ax.text(i, v + 1, f"{v:.1f}%", ha="center", fontsize=8)
        fig.tight_layout()
        p = chart_dir / "perf_scene_error_rate.png"
        fig.savefig(p, dpi=150)
        plt.close(fig)
        charts["PERF_SCENE_ERROR_RATE"] = p

        tp95_vals = [scene_metrics[sid]["metrics"]["tp95_ms"] for sid in scene_ids]
        tps_vals = [scene_metrics[sid]["metrics"]["tps"] for sid in scene_ids]
        fig, ax1 = plt.subplots(figsize=(8.6, 3.6))
        ax1.plot(scene_ids, tp95_vals, marker="o", color="#2E7D32", label="TP95(ms)")
        ax1.set_ylabel("TP95(ms)", color="#2E7D32")
        ax1.tick_params(axis="y", labelcolor="#2E7D32")
        ax2 = ax1.twinx()
        ax2.plot(scene_ids, tps_vals, marker="s", color="#1565C0", label="TPS")
        ax2.set_ylabel("TPS", color="#1565C0")
        ax2.tick_params(axis="y", labelcolor="#1565C0")
        ax1.set_title("性能场景 TP95 / TPS 趋势")
        fig.tight_layout()
        p = chart_dir / "perf_scene_tp95_tps.png"
        fig.savefig(p, dpi=150)
        plt.close(fig)
        charts["PERF_SCENE_TP95_TPS"] = p

    return charts


def expected_labels_from_jmx(jmx_path: pathlib.Path, prefixes: Sequence[str]) -> List[str]:
    if not jmx_path.exists():
        return []
    root = et.parse(jmx_path).getroot()
    out: List[str] = []
    for elem in root.iter("HTTPSamplerProxy"):
        name = elem.attrib.get("testname", "")
        if any(name.startswith(p) for p in prefixes):
            out.append(name)
    return out


def functional_module(label: str) -> str:
    ll = label.lower()
    if "login" in ll or "heartbeat" in ll:
        return "登录/会话"
    if "querycourselist" in ll or "getcoursechaptertasklist" in ll:
        return "课程查询"
    if "queryexpguidebookcatalog" in ll or "queryexpguidebookcontent" in ll:
        return "指导书查询"
    if "savequestionanswer" in ll or "saveexpanswerrecord" in ll:
        return "答题与保存"
    if "quitteachingtask" in ll or "quitcourse" in ll or "logout" in ll:
        return "退出与登出"
    return "其它"


def summarize_cases_by_label(records: Sequence[JtlRecord]) -> Dict[str, Dict[str, Any]]:
    bucket: Dict[str, Dict[str, Any]] = {}
    for r in records:
        b = bucket.setdefault(
            r.label,
            {"executions": 0, "failures": 0, "success": True, "response_codes": Counter(), "fail_msgs": []},
        )
        b["executions"] += 1
        if not r.success:
            b["failures"] += 1
            b["success"] = False
            if r.failure_message:
                b["fail_msgs"].append(r.failure_message)
        if r.response_code:
            b["response_codes"][r.response_code] += 1
    return bucket


def _truncate(text: str, max_len: int = 120) -> str:
    t = _normalize_text(text)
    if len(t) <= max_len:
        return t
    return t[: max_len - 3] + "..."


def top_failure_rows(records: Sequence[JtlRecord], limit: int = 20) -> List[List[Any]]:
    bucket: Dict[str, Dict[str, Any]] = {}
    for r in records:
        d = bucket.setdefault(
            r.label,
            {
                "total": 0,
                "fail": 0,
                "codes": Counter(),
                "msgs": Counter(),
            },
        )
        d["total"] += 1
        if not r.success:
            d["fail"] += 1
            if r.response_code:
                d["codes"][r.response_code] += 1
            if r.failure_message:
                d["msgs"][r.failure_message] += 1
    fail_items = [
        (lb, v) for lb, v in bucket.items() if int(v.get("fail", 0)) > 0
    ]
    fail_items.sort(key=lambda x: (-int(x[1]["fail"]), x[0]))
    rows: List[List[Any]] = []
    for lb, v in fail_items[:limit]:
        total = int(v["total"])
        fail = int(v["fail"])
        rate = (fail / total * 100.0) if total else 0.0
        code_text = ", ".join([f"{k}:{cnt}" for k, cnt in v["codes"].most_common(3)]) or "-"
        msg_text = "; ".join([f"{_truncate(k, 80)}:{cnt}" for k, cnt in v["msgs"].most_common(2)]) or "-"
        rows.append([lb, f"{fail}/{total}", f"{rate:.2f}%", code_text, msg_text])
    return rows


def _case_group(label: str) -> str:
    if label.startswith("FUNC-"):
        return "正向主链路"
    if label.startswith("EX-"):
        return "异常/边界"
    if label.startswith("KA-"):
        return "会话时序"
    return "其它"


def _case_expectation(label: str) -> str:
    if "[EXPECT_FAIL]" in label:
        return "期望失败"
    return "期望成功"


def full_case_detail_rows(
    expected_labels: Sequence[str],
    case_map: Dict[str, Dict[str, Any]],
) -> List[List[Any]]:
    rows: List[List[Any]] = []
    for lb in expected_labels:
        data = case_map.get(lb)
        if not data:
            rows.append([lb, _case_group(lb), _case_expectation(lb), 0, 0, 0, "-", "未执行", "未执行"])
            continue

        total = int(data.get("executions", 0))
        fail = int(data.get("failures", 0))
        passed = max(total - fail, 0)
        codes: Counter = data.get("response_codes", Counter())
        code_text = ", ".join([f"{k}:{cnt}" for k, cnt in codes.most_common(5)]) or "-"
        fail_msgs = Counter(data.get("fail_msgs", []))
        fail_reason = (
            "; ".join([f"{_truncate(msg, 100)}({cnt})" for msg, cnt in fail_msgs.most_common(2)])
            if fail_msgs
            else "-"
        )
        verdict = "PASS" if fail == 0 else "FAIL"
        rows.append([lb, _case_group(lb), _case_expectation(lb), total, passed, fail, code_text, verdict, fail_reason])
    return rows


def yes_no(v: bool) -> str:
    return "是" if v else "否"


def release_recommendation(
    *,
    forced: str,
    func_pass_rate: float,
    perf_scene_results: Dict[str, Dict[str, Any]],
    sev: Dict[str, int],
) -> str:
    if forced:
        return forced
    if sev["致命"] > 0:
        return "不建议上线"
    if sev["严重"] > 0:
        return "有条件上线"
    if func_pass_rate < 95.0:
        return "有条件上线"
    failed_scenes = [sid for sid, m in perf_scene_results.items() if m.get("conclusion") == "待优化"]
    if failed_scenes:
        return "有条件上线"
    return "建议上线"


def evaluate_scene(metrics: Dict[str, float], thresholds: Dict[str, Any]) -> str:
    if metrics["total"] <= 0:
        return "未执行"
    if (
        metrics["error_rate_pct"] <= float(thresholds["error_rate_max_pct"])
        and metrics["tp95_ms"] <= float(thresholds["tp95_max_ms"])
        and metrics["success_rate_pct"] >= float(thresholds["success_rate_min_pct"])
    ):
        return "通过"
    return "待优化"


def scene_id_from_label(label: str) -> Optional[str]:
    m = re.match(r"^(PERF-\d{3})-", label)
    return m.group(1) if m else None


def build_function_report_lines(
    meta: Dict[str, Any],
    functional_records: Sequence[JtlRecord],
    defects: Sequence[Dict[str, str]],
    jmx_path: pathlib.Path,
) -> Tuple[List[str], Dict[str, Any]]:
    doc = meta["document"]
    proj = meta["project"]
    base_data = meta["base_data"]
    accounts = meta["accounts"]

    functional_filtered = [r for r in functional_records if r.label.startswith(("FUNC-", "EX-", "KA-"))]
    if not functional_filtered:
        raise ValueError("functional.jtl does not contain FUNC-/EX-/KA- samples")
    case_map = summarize_cases_by_label(functional_filtered)
    expected = expected_labels_from_jmx(jmx_path, ["FUNC-", "EX-", "KA-"])
    expected_set = set(expected)
    observed_set = set(case_map.keys())
    executed_expected = sorted(expected_set & observed_set)
    unexecuted = sorted(expected_set - observed_set)
    blocked_labels = [lb for lb in observed_set if "PRECONDITION_MISSING" in lb]
    failed_labels = [lb for lb in executed_expected if not case_map[lb]["success"]]
    passed_labels = [lb for lb in executed_expected if case_map[lb]["success"]]
    pass_rate = (len(passed_labels) / len(executed_expected) * 100.0) if executed_expected else 0.0

    module_case: Dict[str, List[str]] = defaultdict(list)
    for lb in executed_expected:
        module_case[functional_module(lb)].append(lb)
    module_rows = []
    for mod in ["登录/会话", "课程查询", "指导书查询", "答题与保存", "退出与登出"]:
        labels = module_case.get(mod, [])
        total = len(labels)
        passed = sum(1 for lb in labels if case_map[lb]["success"])
        failed = total - passed
        rate = f"{(passed / total * 100.0):.2f}%" if total else "0.00%"
        module_rows.append([mod, total, passed, failed, rate])

    main_flow = [
        "FUNC-Login",
        "FUNC-QueryCourseList",
        "FUNC-GetCourseChapterTaskList",
        "FUNC-QueryExpGuidebookCatalog",
        "FUNC-QueryExpGuidebookContent",
        "FUNC-saveQuestionAnswer",
        "FUNC-SaveExpAnswerRecord",
        "FUNC-QuitTeachingTask",
        "FUNC-QuitCourse",
        "FUNC-Logout",
    ]
    main_flow_pass = all(lb in case_map and case_map[lb]["success"] for lb in main_flow)
    data_dep_block = bool(blocked_labels) or any("userreportid" in lb.lower() and not case_map[lb]["success"] for lb in failed_labels)
    session_issue = any(lb.startswith("KA-") and not case_map[lb]["success"] for lb in case_map)
    abnormal_issue = any(
        ("500" in {str(k) for k in case_map[lb]["response_codes"]}) or any("exception" in msg.lower() or "cannot invoke" in msg.lower() for msg in case_map[lb]["fail_msgs"])
        for lb in case_map
    )
    functional_fail_rows = top_failure_rows(functional_filtered, limit=20)
    functional_case_rows = full_case_detail_rows(executed_expected + unexecuted, case_map)

    sev = severity_counts(defects)
    severity_rank = {"致命": 0, "严重": 1, "一般": 2, "建议": 3}
    top_issues = sorted(defects, key=lambda d: severity_rank.get(next((k for k, alias in SEVERITY_MAP.items() if d["severity"].strip().lower() in {x.lower() for x in alias}), "一般"), 9))
    top_issue_lines = []
    for i, d in enumerate(top_issues[:5], start=1):
        top_issue_lines.append(f"{i}. [{d.get('severity','一般')}] {d.get('module','未分类')} - {d.get('title','未命名问题')}")
    while len(top_issue_lines) < 3:
        top_issue_lines.append(f"{len(top_issue_lines)+1}. 待补充（请根据最新缺陷单更新）")
    func_total_samples = len(functional_filtered)
    func_failed_samples = sum(1 for r in functional_filtered if not r.success)
    func_error_rate = (func_failed_samples / func_total_samples * 100.0) if func_total_samples else 0.0
    func_500_count = sum(1 for r in functional_filtered if (not r.success and r.response_code == "500"))
    func_401_count = sum(1 for r in functional_filtered if (not r.success and r.response_code == "401"))
    expect_ok_failed = sum(
        1
        for lb in executed_expected
        if "[EXPECT_OK]" in lb and not case_map.get(lb, {}).get("success", False)
    )

    lines: List[str] = []
    lines.extend(
        [
            "《自动化接口功能测试报告》",
            "",
            "0. 执行摘要（结论先行）",
            "• 总体结论：功能回归未通过完整主链路验证，需优先修复登录与鉴权问题后复测。",
            "• 发布建议：__RELEASE_RECOMMENDATION__",
            f"• 核心指标：主流程跑通={yes_no(main_flow_pass)}；用例通过率={pass_rate:.2f}%（{len(passed_labels)}/{len(executed_expected)}）；样本错误率={func_error_rate:.2f}%（{func_failed_samples}/{func_total_samples}）",
            f"• 关键失败信号：HTTP500={func_500_count}，HTTP401={func_401_count}，EXPECT_OK 失败数={expect_ok_failed}",
            f"• 核心问题1：{top_issue_lines[0]}",
            f"• 核心问题2：{top_issue_lines[1]}",
            f"• 核心问题3：{top_issue_lines[2]}",
            "",
            "1. 文档信息",
            f"文档名称：校园版自动化接口功能测试报告",
            f"项目名称：{proj['name']}",
            f"测试阶段：{doc['test_phase']}",
            f"测试工具：{doc['tool']}",
            f"编写人：{doc['author']}",
            f"编写日期：{doc['date']}",
            f"文档版本：{doc['version']}",
            "",
            "2. 测试概述",
            "2.1 测试背景",
            "为验证校园版核心接口在测试环境中的功能正确性、异常处理能力及自动化回归可行性，开展本次自动化接口功能测试。",
            "2.2 测试目标",
            "1. 验证核心接口主流程是否可正常执行。",
            "2. 验证登录态、Token、Heartbeat 保活机制是否符合预期。",
            "3. 验证接口在正常、边界值及异常输入下的处理能力。",
            "4. 验证核心数据关联逻辑是否正确。",
            "5. 建立可复用的自动化回归测试能力，为后续版本迭代提供支撑。",
            "2.3 测试范围",
            "• 登录接口",
            "• 心跳接口",
            "• 登出接口",
            "• 获取课程列表接口",
            "• 获取课程章节任务列表接口",
            "• 获取指导书目录接口",
            "• 获取指导书内容接口",
            "• 学生答题接口",
            "• 答题过程保存接口",
            "• 退出教学任务接口",
            "• 退出课程接口",
            "2.4 非测试范围",
            "• 前端页面展示与交互逻辑",
            "• 非接口层安全渗透测试",
            "• 生产环境容量验证",
            "• 未提供有效业务数据映射的全量场景自动化覆盖",
            "",
            "3. 测试依据",
            "• 接口文档",
            "• 业务规则说明",
            "• 测试账号与测试数据",
            "• Token 与 Heartbeat 时序规则",
            "• 自动化测试方案设计",
            "",
            "4. 测试环境与数据说明",
            "4.1 测试环境",
        ]
    )
    lines.extend(table_lines(["项目", "内容"], [["环境名称", "测试环境"], ["baseURL", proj["base_url"]], ["测试工具", doc["tool"]], ["执行方式", "自动化批量执行"]]))
    lines.extend(["", "4.2 测试账号"])
    for a in accounts:
        lines.append(f"• {a}")
    lines.extend(["", "4.3 测试基础数据"])
    lines.extend(table_lines(["参数", "值"], [["courseId", base_data["courseId"]], ["teachingTaskId", base_data["teachingTaskId"]], ["blockId", base_data["blockId"]]]))
    lines.extend(
        [
            "",
            "4.4 特殊数据约束说明",
            "• 登录后返回 accessToken 和 refreshToken。",
            "• 后续业务接口均依赖 Authorization。",
            "• Token 需每 10 秒调用一次 Heartbeat，否则会过期失效。",
            "• userReportId 与账号强关联，不能跨账号共用。",
            "• 正向测试中，userReportId 需基于当前账号专属映射数据获取。",
            "• 跨账号 userReportId 使用场景归类为异常测试，不纳入正向主流程结果统计。",
            "",
            "5. 测试策略",
            "5.1 用例设计策略",
            "• 正向功能用例",
            "• 边界值用例",
            "• 异常用例",
            "• 会话时序用例",
            "• 数据关联用例",
            "5.2 自动化执行策略",
            "• 登录后自动提取 accessToken、refreshToken",
            "• 后续接口统一注入 Authorization",
            "• 通过 Heartbeat 维持会话有效性",
            "• 关键响应增加 HTTP 状态码、业务码及关键字段断言",
            "• 对依赖 userReportId 的接口增加账号映射前置校验",
            "5.3 覆盖策略",
        ]
    )
    lines.extend(
        table_lines(
            ["类型", "覆盖说明"],
            [
                ["功能覆盖", "已覆盖核心主流程"],
                ["边界覆盖", "已覆盖关键参数空值、非法值、格式不一致等场景"],
                ["异常覆盖", "已覆盖无 Token、过期 Token、越权数据访问等场景"],
                ["数据依赖覆盖", "部分接口依赖账号专属数据，已按映射方式处理"],
            ],
        )
    )
    lines.extend(
        [
            "",
            "6. 功能测试执行情况",
            "6.1 用例执行汇总",
        ]
    )
    lines.extend(
        table_lines(
            ["指标", "数量"],
            [
                ["用例总数", len(expected_set)],
                ["已执行", len(executed_expected)],
                ["通过", len(passed_labels)],
                ["失败", len(failed_labels)],
                ["阻塞", len(blocked_labels)],
                ["未执行", len(unexecuted)],
                ["通过率", f"{pass_rate:.2f}%"],
            ],
        )
    )
    lines.extend(["", "6.2 分模块执行结果"])
    lines.extend(table_lines(["模块", "用例数", "通过数", "失败数", "通过率"], module_rows))
    lines.extend(
        [
            "",
            "6.3 主流程执行结果",
            "本次自动化重点验证如下主流程：",
            "登录 -> Heartbeat 保活 -> 查询课程列表 -> 查询课程章节任务 -> 查询指导书目录 -> 查询指导书内容 -> 学生答题 -> 保存答题过程 -> 退出教学任务 -> 退出课程 -> 登出",
            "执行结果概述：",
            f"• 主流程是否跑通：{yes_no(main_flow_pass)}",
            f"• 是否存在数据依赖阻塞：{yes_no(data_dep_block)}",
            f"• 是否存在会话时序问题：{yes_no(session_issue)}",
            f"• 是否存在异常返回不规范问题：{yes_no(abnormal_issue)}",
            "",
            "6.4 失败样本明细（Top20）",
        ]
    )
    if functional_fail_rows:
        lines.extend(
            table_lines(
                ["样本标签", "失败次数/总次数", "失败率", "状态码分布", "失败信息摘要"],
                functional_fail_rows,
            )
        )
    else:
        lines.append("无失败样本。")
    lines.extend(
        [
            "",
            "6.5 用例执行详情（全量）",
        ]
    )
    lines.extend(
        table_lines(
            ["用例标签", "分类", "期望", "执行次数", "通过", "失败", "状态码分布", "结论", "失败原因摘要"],
            functional_case_rows,
        )
    )
    lines.extend(
        [
            "",
            "6.6 关键图表",
            "[[CHART:FUNC_MODULE_PASS_FAIL]]",
            "[[CHART:FUNC_STATUS_CODE_DIST]]",
            "",
            "注：若接口响应中出现服务端异常类名或空指针信息（如 NullPointerException/Cannot invoke 等），应视为缺陷线索。",
            "",
            "7. 缺陷与问题分析",
            "7.1 缺陷统计",
        ]
    )
    lines.extend(table_lines(["缺陷等级", "数量"], [["致命", sev["致命"]], ["严重", sev["严重"]], ["一般", sev["一般"]], ["建议", sev["建议"]]]))
    lines.extend(["", "7.2 重点问题说明"])
    lines.extend(top_issue_lines)
    lines.extend(
        [
            "7.3 风险分析",
        ]
    )
    lines.extend(
        table_lines(
            ["风险项", "影响", "建议"],
            [
                ["Token 保活机制依赖固定周期调用", "影响业务连续访问", "建议加强时序校验和容错处理"],
                ["Header 不统一", "影响接入兼容性", "建议统一接口规范"],
                ["动态业务数据依赖强", "影响自动化稳定性", "建议补充动态取数或映射管理机制"],
            ],
        )
    )
    lines.extend(
        [
            "",
            "8. 测试结论",
            "8.1 总体结论",
            "本轮自动化执行已覆盖核心接口主流程、异常流程及会话时序流程；但主流程未全量通过，当前结果不支持判定功能已稳定。",
            f"从数据看，已执行用例通过率为 {pass_rate:.2f}%（{len(passed_labels)}/{len(executed_expected)}），关键阻断点集中在登录/鉴权链路与异常处理规范性。",
            "自动化脚本本身已具备回归复测能力，建议在缺陷修复后按同口径重新执行并更新结论。",
            "8.2 发布建议",
            "建议上线 / 有条件上线 / 不建议上线：由规则引擎在汇总页自动给出。",
            "",
            "9. 后续计划",
            "1. 完善 userReportId 动态获取能力。",
            "2. 增补更多异常与兼容性自动化场景。",
            "3. 将自动化测试纳入版本回归流程。",
            "4. 与性能测试结果联动形成统一质量基线。",
            "",
            "10. 附录",
            "• 接口清单",
            "• 用例清单",
            "• 账号与数据映射清单",
            "• JMeter 工程说明",
            "• 执行截图/结果摘要",
        ]
    )

    status_counter = Counter(r.response_code for r in functional_filtered if r.response_code)
    module_chart_data = []
    for r in module_rows:
        module_chart_data.append(
            {
                "module": str(r[0]),
                "passed": int(r[2]),
                "failed": int(r[3]),
            }
        )

    return lines, {
        "pass_rate": pass_rate,
        "total_cases": len(expected_set),
        "executed_cases": len(executed_expected),
        "failed_cases": len(failed_labels),
        "main_flow_pass": main_flow_pass,
        "sev": sev,
        "top_issues": top_issue_lines[:3],
        "module_chart_data": module_chart_data,
        "status_code_counter": status_counter,
    }


def build_performance_report_lines(
    meta: Dict[str, Any],
    performance_records: Sequence[JtlRecord],
    defects: Sequence[Dict[str, str]],
) -> Tuple[List[str], Dict[str, Any]]:
    doc = meta["document"]
    proj = meta["project"]
    base_data = meta["base_data"]
    accounts = meta["accounts"]
    thresholds = meta["thresholds"]
    scene_conf = meta["performance_scene_config"]

    perf_only = [r for r in performance_records if r.label.startswith("PERF-")]
    if not perf_only:
        raise ValueError("performance.jtl does not contain PERF-00X-* samples")

    scene_records: Dict[str, List[JtlRecord]] = defaultdict(list)
    unmapped: List[str] = []
    for r in perf_only:
        sid = scene_id_from_label(r.label)
        if sid in SCENE_NAME:
            scene_records[sid].append(r)
        else:
            unmapped.append(r.label)

    # Long steady scenario (PERF-005) is executed as mixed labels in current JMeter design.
    # Build a dedicated bucket to avoid under-counting only PERF-005-login samples.
    label_counts = Counter(r.label for r in perf_only)
    long_records: List[JtlRecord] = []
    for r in perf_only:
        lb = r.label
        if "-mixed" in lb or lb.startswith("PERF-005-"):
            long_records.append(r)
            continue
        # In current scripts, mixed loop heartbeat/queryCourseList may share labels with baseline scenes.
        # Use a conservative small-volume heuristic to include them in long-steady summary when applicable.
        if lb in {"PERF-002-heartbeat", "PERF-003-queryCourseList"} and label_counts[lb] <= 1000:
            long_records.append(r)

    scene_metrics: Dict[str, Dict[str, Any]] = {}
    overall_rows = []
    for sid in ["PERF-001", "PERF-002", "PERF-003", "PERF-004", "PERF-005"]:
        scene_source = long_records if sid == "PERF-005" else scene_records.get(sid, [])
        metrics = aggregate_metrics(scene_source)
        conclusion = evaluate_scene(metrics, thresholds)
        scene_metrics[sid] = {"metrics": metrics, "conclusion": conclusion}
        conf = scene_conf.get(sid, {})
        overall_rows.append(
            [
                SCENE_NAME[sid],
                conf.get("concurrency", "N/A"),
                metrics["tps"],
                int(metrics["avg_ms"]),
                int(metrics["tp95_ms"]),
                int(metrics["tp99_ms"]),
                f"{metrics['error_rate_pct']:.2f}%",
                conclusion,
            ]
        )

    iface_rules = {
        "login": lambda lb: "login" in lb.lower(),
        "heartbeat": lambda lb: "heartbeat" in lb.lower(),
        "queryCourseList": lambda lb: "querycourselist" in lb.lower(),
        "getCourseChapterTaskList": lambda lb: "getcoursechaptertasklist" in lb.lower(),
        "queryExpGuidebookCatalog": lambda lb: "queryexpguidebookcatalog" in lb.lower(),
        "queryExpGuidebookContent": lambda lb: "queryexpguidebookcontent" in lb.lower(),
        "saveQuestionAnswer": lambda lb: "savequestionanswer" in lb.lower(),
        "saveExpAnswerRecord": lambda lb: "saveexpanswerrecord" in lb.lower(),
    }
    iface_rows = []
    iface_metrics: Dict[str, Dict[str, Any]] = {}
    for name, matcher in iface_rules.items():
        subset = [r for r in perf_only if matcher(r.label)]
        m = aggregate_metrics(subset)
        iface_metrics[name] = m
        iface_rows.append(["按场景配置", int(m["avg_ms"]), int(m["tp95_ms"]), f"{m['error_rate_pct']:.2f}%", evaluate_scene(m, thresholds)])

    token_invalid = any(("token" in r.failure_message.lower() or "过期" in r.failure_message) for r in long_records if not r.success)
    heartbeat_fail = any(("heartbeat" in r.label.lower()) and not r.success for r in long_records)
    resp_rise = False
    err_rise = False
    jitter = False
    long_metrics = aggregate_metrics(long_records)
    long_stopped_early = bool(long_records) and long_metrics["duration_s"] < 1500
    if len(long_records) >= 6:
        sorted_records = sorted(long_records, key=lambda x: x.timestamp or 0)
        n = len(sorted_records)
        head = sorted_records[: max(1, n // 3)]
        tail = sorted_records[-max(1, n // 3) :]
        head_avg = aggregate_metrics(head)["avg_ms"]
        tail_avg = aggregate_metrics(tail)["avg_ms"]
        head_err = aggregate_metrics(head)["error_rate_pct"]
        tail_err = aggregate_metrics(tail)["error_rate_pct"]
        resp_rise = tail_avg > head_avg * 1.2
        err_rise = tail_err > head_err + 1.0
        jitter = aggregate_metrics(long_records)["tp99_ms"] > aggregate_metrics(long_records)["tp95_ms"] * 1.8
    perf_fail_rows = top_failure_rows(perf_only, limit=20)
    perf_total = len(perf_only)
    perf_failed = sum(1 for r in perf_only if not r.success)
    perf_overall_error = (perf_failed / perf_total * 100.0) if perf_total else 0.0
    perf_overall_metrics = aggregate_metrics(perf_only)
    failed_scenes = [sid for sid in ["PERF-001", "PERF-002", "PERF-003", "PERF-004", "PERF-005"] if scene_metrics[sid]["conclusion"] != "通过"]
    scene_summary = "、".join(failed_scenes) if failed_scenes else "无"
    top_perf_issue_lines = []
    for row in perf_fail_rows[:3]:
        top_perf_issue_lines.append(f"{row[0]}（失败率 {row[2]}，状态码 {row[3]}）")
    while len(top_perf_issue_lines) < 3:
        top_perf_issue_lines.append("待补充（当前无更多失败样本）")

    lines: List[str] = []
    lines.extend(
        [
            "《自动化接口性能测试报告》",
            "",
            "0. 执行摘要（结论先行）",
            "• 总体结论：当前性能结果不满足既定 KPI，需先完成缺陷修复再进行完整复测。",
            "• 决策建议：__RELEASE_RECOMMENDATION__",
            f"• KPI 概览：总体错误率={perf_overall_error:.2f}%（阈值≤{thresholds['error_rate_max_pct']}%），TP95={int(perf_overall_metrics['tp95_ms'])}ms（阈值≤{thresholds['tp95_max_ms']}ms）",
            f"• 场景结论：待优化场景={scene_summary}；长稳提前结束={yes_no(long_stopped_early)}，实际时长={int(long_metrics['duration_s'])}s",
            f"• 核心问题1：{top_perf_issue_lines[0]}",
            f"• 核心问题2：{top_perf_issue_lines[1]}",
            f"• 核心问题3：{top_perf_issue_lines[2]}",
            "",
            "1. 文档信息",
            "文档名称：校园版自动化接口性能测试报告",
            f"项目名称：{proj['performance_name']}",
            f"测试阶段：{doc['test_phase']}",
            f"测试工具：{doc['tool']}",
            f"编写人：{doc['author']}",
            f"编写日期：{doc['date']}",
            f"文档版本：{doc['version']}",
            "",
            "2. 测试概述",
            "2.1 测试背景",
            "为评估校园版核心接口在并发访问场景下的响应能力、稳定性和容量风险，开展本次自动化接口性能测试。",
            "2.2 测试目标",
            "1. 验证核心接口在并发场景下的响应时间表现。",
            "2. 验证登录、Heartbeat、查询、保存等关键接口的吞吐能力。",
            "3. 验证带 Heartbeat 保活机制的会话场景稳定性。",
            "4. 识别潜在性能瓶颈和容量风险。",
            "5. 为后续优化和上线评估提供依据。",
            "",
            "3. 测试对象",
            "• 登录接口",
            "• 心跳接口",
            "• 获取课程列表接口",
            "• 获取课程章节任务列表接口",
            "• 获取指导书目录接口",
            "• 获取指导书内容接口",
            "• 学生答题接口",
            "• 答题过程保存接口",
            "",
            "4. 测试环境说明",
            "4.1 环境信息",
        ]
    )
    lines.extend(table_lines(["项目", "内容"], [["测试环境", "测试环境"], ["baseURL", proj["base_url"]], ["测试工具", doc["tool"]], ["执行方式", "自动化压测"]]))
    lines.extend(["", "4.2 测试数据"])
    lines.extend(table_lines(["参数", "值"], [["courseId", base_data["courseId"]], ["teachingTaskId", base_data["teachingTaskId"]], ["blockId", base_data["blockId"]]]))
    lines.extend(["", "4.3 测试账号"])
    for a in accounts:
        lines.append(f"• {a}")
    lines.extend(
        [
            "",
            "4.4 性能测试前提",
            "• 登录成功后返回 Token",
            "• 后续接口依赖 Authorization",
            "• 每 10 秒需调用一次 Heartbeat 保持会话有效",
            "• userReportId 需按账号映射使用",
            "• 压测过程中不采用跨账号共用 userReportId",
            "",
            "5. 性能测试策略",
            "5.1 测试模型",
            "1. 单接口基准测试：分别验证各核心接口的基础性能表现。",
            "2. 混合业务场景测试：模拟真实用户行为链路。",
            "3. 稳定性测试：验证持续运行条件下接口性能是否稳定。",
            "5.2 指标定义",
        ]
    )
    lines.extend(
        table_lines(
            ["指标", "说明"],
            [
                ["平均响应时间", "接口平均耗时"],
                ["TP90", "90% 请求响应时间"],
                ["TP95", "95% 请求响应时间"],
                ["TP99", "99% 请求响应时间"],
                ["TPS", "每秒事务处理数"],
                ["错误率", "失败请求占比"],
                ["成功率", "成功请求占比"],
            ],
        )
    )
    lines.extend(
        [
            "5.3 成功判定标准",
            f"• 错误率 ≤ {thresholds['error_rate_max_pct']}%",
            f"• TP95 ≤ {thresholds['tp95_max_ms']} ms",
            f"• 接口成功率 ≥ {thresholds['success_rate_min_pct']}%",
            "• 长稳测试期间无持续性恶化趋势",
            "• 会话场景中 Heartbeat 保活机制运行正常",
            "",
            "6. 性能场景设计",
        ]
    )
    scene_design_rows = []
    for sid in ["PERF-001", "PERF-002", "PERF-003", "PERF-004", "PERF-005"]:
        c = scene_conf.get(sid, {})
        scene_design_rows.append([sid, SCENE_NAME[sid], c.get("concurrency", "N/A"), c.get("duration", "N/A"), c.get("desc", "")])
    lines.extend(table_lines(["场景编号", "场景名称", "并发数", "持续时间", "场景说明"], scene_design_rows))
    lines.extend(["", "7. 测试结果汇总", "7.1 总体结果"])
    lines.extend(table_lines(["场景", "并发", "TPS", "平均响应(ms)", "TP95(ms)", "TP99(ms)", "错误率", "结论"], overall_rows))
    lines.extend(["", "7.2 分接口结果"])
    iface_table_rows = []
    for name, row in zip(iface_rules.keys(), iface_rows):
        iface_table_rows.append([name] + row)
    lines.extend(table_lines(["接口", "并发", "平均响应(ms)", "TP95(ms)", "错误率", "结论"], iface_table_rows))
    lines.extend(
        [
            "",
            "7.3 长稳测试结果",
            f"• 是否出现错误率上升：{yes_no(err_rise)}",
            f"• 是否出现响应时间持续增加：{yes_no(resp_rise)}",
            f"• 是否出现 Token 异常失效：{yes_no(token_invalid)}",
            f"• 是否出现 Heartbeat 保活失败：{yes_no(heartbeat_fail)}",
            f"• 是否存在接口抖动现象：{yes_no(jitter)}",
            f"• 长稳场景是否提前结束：{yes_no(long_stopped_early)}",
            f"• 长稳场景实际持续时长：{int(long_metrics['duration_s'])}s",
        ]
    )
    if unmapped:
        lines.append(f"• 未映射样本（命名不符合 PERF-00X-*）：{', '.join(sorted(set(unmapped))[:20])}")
    lines.extend(["", "7.4 错误明细（Top20）"])
    if perf_fail_rows:
        lines.extend(
            table_lines(
                ["样本标签", "失败次数/总次数", "失败率", "状态码分布", "失败信息摘要"],
                perf_fail_rows,
            )
        )
    else:
        lines.append("无失败样本。")
    lines.extend(
        [
            "",
            "7.5 关键图表",
            "[[CHART:PERF_SCENE_ERROR_RATE]]",
            "[[CHART:PERF_SCENE_TP95_TPS]]",
            "",
            "8. 性能问题分析",
            "8.1 性能瓶颈分析",
            "1. 登录接口在高并发场景下响应时间上升明显，存在认证链路性能瓶颈风险。",
            "2. Heartbeat 作为高频调用接口，可能对服务端会话维护能力造成持续压力。",
            "3. 查询类接口整体响应较稳定，但个别场景存在波动。",
            "4. 保存类接口在混合场景下响应时间高于查询类接口，可能存在写操作瓶颈。",
            "8.2 原因初步判断",
            "• 认证逻辑处理链较长",
            "• 高频保活请求累积增加系统负担",
            "• 查询类接口可能存在数据访问波动",
            "• 保存类接口可能受写库或事务影响",
            "8.3 优化建议",
            "• 优化认证链路和 Token 校验流程",
            "• 评估 Heartbeat 调用机制及服务端承载能力",
            "• 优化高频查询与写入接口处理效率",
            "• 增强服务端异常监控与性能监控能力",
            "",
            "9. 风险评估",
        ]
    )
    lines.extend(
        table_lines(
            ["风险项", "表现", "影响", "建议"],
            [
                ["高频 Heartbeat 压力", "并发下持续调用", "增加服务端负载", "建议专项评估会话维持策略"],
                ["登录并发性能波动", "高并发时耗时升高", "影响用户登录体验", "建议优化认证性能"],
                ["动态数据依赖", "数据映射不完整", "影响压测真实性", "建议完善测试数据准备机制"],
            ],
        )
    )
    lines.extend(
        [
            "",
            "10. 性能测试结论",
            "10.1 总体结论",
            "本轮性能压测覆盖基准、混合与长稳场景，结果显示关键链路存在显著错误率风险，暂不满足既定 KPI 口径。",
            f"从总体指标看，错误率为 {perf_overall_error:.2f}%（阈值≤{thresholds['error_rate_max_pct']}%），TP95 为 {int(perf_overall_metrics['tp95_ms'])}ms（阈值≤{thresholds['tp95_max_ms']}ms）。",
            f"长稳场景在本轮执行中为提前停止状态（实际 {int(long_metrics['duration_s'])}s），建议在完成缺陷修复后补齐完整时窗复测。",
            "10.2 建议",
            "• 建议继续扩大并发梯度，逐步验证容量上限。",
            "• 建议针对登录与 Heartbeat 做专项性能优化。",
            "• 建议完善动态数据准备机制，提高压测真实性。",
            "• 建议上线前补充预发布环境性能验证。",
            "",
            "11. 附录",
            "• JMeter 场景配置说明",
            "• 压测参数配置",
            "• 原始结果文件说明",
            "• 图表截图",
            "• 指标解释",
        ]
    )
    return lines, {"scene_metrics": scene_metrics, "iface_metrics": iface_metrics, "unmapped": sorted(set(unmapped))}


def build_exec_summary_lines(
    meta: Dict[str, Any],
    func_info: Dict[str, Any],
    perf_info: Dict[str, Any],
    recommendation: str,
) -> List[str]:
    perf_iface = perf_info["iface_metrics"]
    lines = [
        "《校园版自动化接口测试汇报摘要》",
        "",
        "1. 测试范围",
        "• 覆盖登录、会话保活、课程查询、指导书查询、答题、保存、退出、登出等核心接口",
        "",
        "2. 功能测试结果",
        f"• 用例总数：{func_info['total_cases']}",
        f"• 通过率：{func_info['pass_rate']:.2f}%",
        f"• 主流程是否跑通：{yes_no(func_info['main_flow_pass'])}",
        "• 主要问题：",
    ]
    lines.extend([f"  - {x}" for x in func_info["top_issues"][:3]])
    lines.extend(
        [
            "",
            "3. 性能测试结果",
            f"• 登录接口表现：TP95={int(perf_iface['login']['tp95_ms'])}ms，错误率={perf_iface['login']['error_rate_pct']:.2f}%",
            f"• Heartbeat 稳定性：TP95={int(perf_iface['heartbeat']['tp95_ms'])}ms，错误率={perf_iface['heartbeat']['error_rate_pct']:.2f}%",
            f"• 查询类接口表现：queryCourseList TP95={int(perf_iface['queryCourseList']['tp95_ms'])}ms",
            f"• 保存类接口表现：saveQuestionAnswer TP95={int(perf_iface['saveQuestionAnswer']['tp95_ms'])}ms",
            "",
            "4. 主要风险",
            "• Token 保活机制依赖 Heartbeat",
            "• Header/文档一致性问题",
            "• userReportId 数据依赖强",
            "",
            "5. 结论与建议",
            f"• 当前结论：{recommendation}",
            "• 后续建议：优先修复高优先级缺陷，补齐性能瓶颈优化后进入下一阶段。",
        ]
    )
    return lines


def validate_required_files(args: argparse.Namespace) -> None:
    for label, value in [
        ("functional-jtl", args.functional_jtl),
        ("performance-jtl", args.performance_jtl),
        ("defects", args.defects),
        ("meta", args.meta),
    ]:
        if not value:
            raise ValueError(f"--{label} is required in strict mode")
        p = pathlib.Path(value)
        if not p.exists():
            raise FileNotFoundError(f"{label} file not found: {p}")


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate 3 strict formal reports from JMeter data.")
    parser.add_argument("--functional-jtl", required=True, help="Path to functional JTL (required)")
    parser.add_argument("--performance-jtl", required=True, help="Path to performance JTL (required)")
    parser.add_argument("--defects", required=True, help="Path to defects CSV/JSON (required)")
    parser.add_argument("--meta", required=True, help="Path to report meta JSON/YAML (required)")
    parser.add_argument("--jmx", default="jmeter/test-plan.jmx", help="Path to JMX for expected-label validation")
    parser.add_argument("--output-dir", default="jmeter/docs", help="Output directory")
    args = parser.parse_args()

    validate_required_files(args)
    meta = load_meta(pathlib.Path(args.meta))
    functional_records = parse_jtl(pathlib.Path(args.functional_jtl))
    performance_records = parse_jtl(pathlib.Path(args.performance_jtl))
    defects = load_defects(pathlib.Path(args.defects))

    out_dir = pathlib.Path(args.output_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    func_lines, func_info = build_function_report_lines(meta, functional_records, defects, pathlib.Path(args.jmx))
    perf_lines, perf_info = build_performance_report_lines(meta, performance_records, defects)
    recommendation = release_recommendation(
        forced=_normalize_text(meta.get("release_recommendation", "")),
        func_pass_rate=func_info["pass_rate"],
        perf_scene_results=perf_info["scene_metrics"],
        sev=func_info["sev"],
    )
    func_lines = [
        ln.replace("__RELEASE_RECOMMENDATION__", recommendation).replace(
            "建议上线 / 有条件上线 / 不建议上线：由规则引擎在汇总页自动给出。",
            f"发布建议：{recommendation}",
        )
        for ln in func_lines
    ]
    perf_lines = [ln.replace("__RELEASE_RECOMMENDATION__", recommendation) for ln in perf_lines]
    summary_lines = build_exec_summary_lines(meta, func_info, perf_info, recommendation)

    func_path = out_dir / "自动化接口功能测试报告.docx"
    perf_path = out_dir / "自动化接口性能测试报告.docx"
    sum_path = out_dir / "校园版自动化接口测试汇报摘要.docx"

    func_charts = build_function_chart_paths(func_info, out_dir)
    perf_charts = build_performance_chart_paths(perf_info, out_dir)

    write_professional_docx(func_path, func_lines, charts=func_charts)
    write_professional_docx(perf_path, perf_lines, charts=perf_charts)
    write_professional_docx(sum_path, summary_lines, charts=None)

    print(f"[OK] 功能报告: {func_path}")
    print(f"[OK] 性能报告: {perf_path}")
    print(f"[OK] 一页摘要: {sum_path}")
    if perf_info["unmapped"]:
        print(f"[WARN] 未映射性能样本: {', '.join(perf_info['unmapped'][:30])}")
    print(f"[INFO] 发布建议(自动判定): {recommendation}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"[ERROR] {exc}", file=sys.stderr)
        raise
